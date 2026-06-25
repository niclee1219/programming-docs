"""
AGM Generator — DOCX Generator
Recreates the Singapore AGM Directors' Resolution document format.
Uses python-docx; faithfully mirrors the formatting of the sample template.
Section content is driven by a profile dict for per-company customisation.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import ORDINALS
import re

# Default profile — matches the original hardcoded document exactly.
# Used when no per-company profile has been saved.
DEFAULT_PROFILE = {
    "banner_pursuant": "Article 93 of the Company’s Articles of Association",
    "sections": [
        {
            "id": "financial_statements",
            "included": True,
            "heading": "AUDITED ACCOUNTS/FINANCIAL STATEMENTS",
            "body": (
                "That the audited accounts/financial statements of the Company for the year ended "
                "{year_end} having been examined by the Directors be and are hereby approved "
                "and authorised for issue."
            ),
        },
        {
            "id": "directors_statement",
            "included": True,
            "heading": "DIRECTORS’ STATEMENT & AUDITORS’ REPORT",
            "body": (
                "That the Directors’ Statement and Report of Auditors of the Company for the "
                "year ended {year_end} be and are hereby adopted and approved and that any two "
                "Directors be authorised to sign on behalf of the Board the Directors’ Statement."
            ),
        },
        {
            "id": "agm_notice",
            "included": True,
            "heading": "{ordinal_upper} ANNUAL GENERAL MEETING",
            "body": (
                "That the {ordinal_cap} Annual General Meeting of members of the Company will be "
                "held at {location} on {agm_date}."
            ),
        },
    ],
    "agenda_items": [
        {
            "id": "a_financial",
            "included": True,
            "text": (
                "To receive and approve the Company’s Audited Accounts/financial statements "
                "together with the Directors’ Statement and Report of Auditors for the year ended {year_end}."
            ),
        },
        {"id": "b_auditors_remuneration", "included": True, "text": "To approve Auditors’ Remuneration"},
        {"id": "c_elect_directors",       "included": True, "text": "To elect Directors"},
        {"id": "d_appoint_auditors",      "included": True, "text": "To appoint Auditors"},
        {"id": "e_any_other",             "included": True, "text": "To transact any other business, if any."},
    ],
}

SMALL_EXEMPT_PROFILE = {
    "banner_pursuant": "Article 93 of the Company’s Articles of Association",
    "sections": [
        {
            "id": "small_exempt_header",
            "included": True,
            "heading": "SMALL COMPANY EXEMPT FROM AUDIT REQUIREMENTS",
            "body": (
                "That the Company being a small company as defined under Section 205C of the Companies Act "
                "is exempt from the requirement to have its financial statements audited."
            ),
        },
        {
            "id": "financial_statements",
            "included": True,
            "heading": "UNAUDITED FINANCIAL STATEMENTS",
            "body": (
                "That the unaudited financial statements of the Company for the year ended "
                "{year_end} having been examined by the Directors be and are hereby approved "
                "and authorised for issue."
            ),
        },
        {
            "id": "directors_statement",
            "included": True,
            "heading": "DIRECTORS’ STATEMENT",
            "body": (
                "That the Directors’ Statement of the Company for the year ended {year_end} "
                "be and is hereby adopted and approved and that any two Directors be authorised "
                "to sign on behalf of the Board the Directors’ Statement."
            ),
        },
        {
            "id": "agm_notice",
            "included": True,
            "heading": "{ordinal_upper} ANNUAL GENERAL MEETING",
            "body": (
                "That the {ordinal_cap} Annual General Meeting of members of the Company will be "
                "held at {location} on {agm_date}."
            ),
        },
    ],
    "agenda_items": [
        {
            "id": "a_financial",
            "included": True,
            "text": (
                "To receive and approve the Company’s Unaudited Financial Statements together "
                "with the Directors’ Statement for the year ended {year_end}."
            ),
        },
        {"id": "c_elect_directors", "included": True, "text": "To elect Directors"},
        {"id": "e_any_other",       "included": True, "text": "To transact any other business, if any."},
    ],
}

# AGM Minutes — structural skeleton only. No default sections yet; content to
# be filled in in a future iteration. Kept in the same shape (sections /
# agenda_items) as the resolution profiles so the existing profile editor and
# preview renderer can be reused unchanged.
MINUTES_DEFAULT_PROFILE = {
    "location": "",
    "sections": [],
    "agenda_items": [],
}


def _add_run(para, text: str, bold=False, size_pt: float = None, underline=False, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def _set_para_spacing(para, before_pt=0, after_pt=0, line_pt=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_pt:
        pf.line_spacing = Pt(line_pt)


def _add_double_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "double")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "5")
        el.set(qn("w:color"), "auto")
        pBdr.append(el)
    pPr.append(pBdr)


class AGMGenerator:

    def generate_to_bytes(self, data: dict, profile: dict = None, doc_type: str = "directors_resolution") -> tuple[bool, bytes, str]:
        try:
            import io
            doc = self._build(doc_type, data, profile)
            buf = io.BytesIO()
            doc.save(buf)
            return True, buf.getvalue(), ""
        except Exception as exc:
            import traceback
            return False, b"", f"Error: {exc}\n{traceback.format_exc()}"

    def generate(self, data: dict, output_path: str, profile: dict = None, doc_type: str = "directors_resolution") -> tuple[bool, str]:
        try:
            doc = self._build(doc_type, data, profile)
            doc.save(output_path)
            from pathlib import Path
            pdf_path = str(Path(output_path).with_suffix(".pdf"))
            pdf_ok = self.convert_docx_to_pdf(output_path, pdf_path)
            return True, f"Saved: {output_path}" + (" (and PDF)" if pdf_ok else " (PDF conversion failed)")
        except Exception as exc:
            import traceback
            return False, f"Error: {exc}\n{traceback.format_exc()}"

    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        import sys, subprocess
        from pathlib import Path
        docx_path_abs = str(Path(docx_path).resolve())
        pdf_path_abs  = str(Path(pdf_path).resolve())

        if sys.platform == "darwin":
            applescript = f'''
            tell application "Microsoft Word"
                set original_setting to screen updating
                set screen updating to false
                try
                    open "{docx_path_abs}"
                    set theDoc to active document
                    save as theDoc file format format pdf file name "{pdf_path_abs}"
                    close theDoc saving no
                on error errText number errNum
                    set screen updating to original_setting
                    error errText number errNum
                end try
                set screen updating to original_setting
            end tell
            '''
            try:
                subprocess.run(["osascript", "-e", applescript], capture_output=True, text=True, check=True)
                return True
            except Exception:
                return False
        elif sys.platform == "win32":
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(docx_path_abs)
                doc.SaveAs(pdf_path_abs, FileFormat=17)
                doc.Close(); word.Quit()
                return True
            except Exception:
                try:
                    import comtypes.client
                    word = comtypes.client.CreateObject("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(docx_path_abs)
                    doc.SaveAs(pdf_path_abs, FileFormat=17)
                    doc.Close(); word.Quit()
                    return True
                except Exception:
                    return False
        else:
            try:
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(Path(pdf_path).parent), docx_path_abs,
                ], check=True, capture_output=True)
                return True
            except Exception:
                return False

    # ------------------------------------------------------------------
    # Document construction
    # ------------------------------------------------------------------

    def _build(self, doc_type: str, data: dict, profile: dict = None) -> Document:
        if doc_type == "agm_minutes":
            return self._build_minutes_document(data, profile)
        return self._build_document(data, profile)

    def _new_document(self) -> Document:
        doc = Document()
        for section in doc.sections:
            section.page_width    = Cm(21)
            section.page_height   = Cm(29.7)
            section.left_margin   = Cm(3.17)
            section.right_margin  = Cm(3.17)
            section.top_margin    = Cm(2.54)
            section.bottom_margin = Cm(2.54)
        return doc

    def _add_company_header(self, doc: Document, data: dict):
        company_name = data.get("company_name", "").strip().upper()
        reg_no       = data.get("reg_no", "").strip()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, after_pt=2)
        _add_run(p, company_name, bold=True, size_pt=14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, after_pt=0)
        _add_run(p, "." * 75, size_pt=10)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, after_pt=0)
        _add_run(p, "Company Regn No.  ", size_pt=9)
        _add_run(p, reg_no, size_pt=9)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, after_pt=4)
        _add_run(p, "(Incorporated in the Republic of Singapore)", size_pt=10)

        p = doc.add_paragraph()
        _set_para_spacing(p, after_pt=0)

    def _build_minutes_document(self, data: dict, profile: dict = None) -> Document:
        """AGM Minutes — structural skeleton. Renders the company letterhead
        and banner, then loops profile sections/agenda items exactly like the
        resolution builder so future content slots in without further wiring.
        """
        doc = self._new_document()
        self._add_company_header(doc, data)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before_pt=4, after_pt=8)
        _add_run(p, "MINUTES OF ANNUAL GENERAL MEETING", bold=True, size_pt=13)

        subs = self._build_subs(data, profile or MINUTES_DEFAULT_PROFILE)
        p_obj = profile or MINUTES_DEFAULT_PROFILE
        sections     = p_obj.get("sections", [])
        agenda_items = p_obj.get("agenda_items", [])

        for sec in sections:
            if not sec.get("included", True):
                continue
            self._section_heading(doc, self._sub(sec.get("heading", ""), subs))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_para_spacing(p, after_pt=6)
            _add_run(p, self._sub(sec.get("body", ""), subs), size_pt=9.5)

        included_items = [it for it in agenda_items if it.get("included", True)]
        if included_items:
            self._section_heading(doc, "AGENDA OF ANNUAL GENERAL MEETING")
            for i, item in enumerate(included_items):
                letter = chr(ord('a') + i)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.first_line_indent = Cm(-1)
                _set_para_spacing(p, after_pt=2)
                _add_run(p, f"({letter})\t", size_pt=9.5)
                _add_run(p, self._sub(item.get("text", ""), subs), size_pt=9.5)

        return doc

    def _build_subs(self, data: dict, p_obj: dict) -> dict:
        address  = data.get("address", "").strip()
        year_end = data.get("financial_year_end", "").strip()
        agm_date = data.get("agm_date", "").strip()

        try:
            agm_num = int(re.sub(r"\D", "", str(data.get("agm_number", "1"))) or 1)
        except ValueError:
            agm_num = 1
        ordinal_cap   = ORDINALS.get(agm_num, f"{agm_num}th")
        ordinal_upper = ordinal_cap.upper()

        return {
            "year_end":      year_end,
            "ordinal_cap":   ordinal_cap,
            "ordinal_upper": ordinal_upper,
            "address":       address,
            "agm_date":      agm_date,
            "location":      p_obj.get("location") or address,
        }

    def _build_document(self, data: dict, profile: dict = None) -> Document:
        doc = self._new_document()

        # ---- Prepare data ----
        directors = [d.strip() for d in data.get("directors", []) if d.strip()]

        p_obj = profile or DEFAULT_PROFILE
        sections     = p_obj.get("sections",      DEFAULT_PROFILE["sections"])
        agenda_items = p_obj.get("agenda_items",  DEFAULT_PROFILE["agenda_items"])
        pursuant     = p_obj.get("banner_pursuant", DEFAULT_PROFILE["banner_pursuant"])
        subs         = self._build_subs(data, p_obj)
        ordinal_cap   = subs["ordinal_cap"]
        ordinal_upper = subs["ordinal_upper"]

        # ---- Company header ----
        self._add_company_header(doc, data)

        # ---- DIRECTORS' RESOLUTION banner ----
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before_pt=4, after_pt=4)
        _add_double_border(p)
        _add_run(p, "DIRECTORS’ RESOLUTION", bold=True, size_pt=30)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before_pt=6, after_pt=0)
        _add_run(p, f"In writing pursuant to the authority given by {pursuant},", size_pt=9.5)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, after_pt=8)
        _add_run(p, "hereby RESOLVED:-", size_pt=9.5)

        # ---- Sections ----
        for sec in sections:
            if not sec.get("included", True):
                continue
            heading = self._sub(sec.get("heading", ""), subs)
            body    = self._sub(sec.get("body", ""), subs)
            self._section_heading(doc, heading)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_para_spacing(p, after_pt=6)
            _add_run(p, body, size_pt=9.5)

        # ---- Agenda ----
        included_items = [it for it in agenda_items if it.get("included", True)]
        if included_items:
            self._section_heading(doc, "AGENDA OF ANNUAL GENERAL MEETING")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_para_spacing(p, after_pt=2)
            _add_run(
                p,
                "That the Agenda of the Annual General Meeting of the members of "
                "the Company be respectively as follows :-",
                size_pt=9.5,
            )
            for i, item in enumerate(included_items):
                letter = chr(ord('a') + i)
                text = self._sub(item.get("text", ""), subs)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.first_line_indent = Cm(-1)
                _set_para_spacing(p, after_pt=2)
                _add_run(p, f"({letter})\t", size_pt=9.5)
                _add_run(p, text, size_pt=9.5)

        # ---- Director signatures ----
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=10, after_pt=2)
        _add_run(p, "DIRECTORS  ", bold=True, size_pt=9.5)

        self._add_director_signatures(doc, directors)

        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=10, after_pt=0)
        _add_run(p, "Dated this                   day of", size_pt=9.5)

        return doc

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _sub(text: str, subs: dict) -> str:
        try:
            return text.format_map(subs)
        except (KeyError, ValueError):
            return text

    def _section_heading(self, doc: Document, text: str):
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=8, after_pt=2)
        _add_run(p, text, bold=True, underline=True, size_pt=9.5)

    def _add_director_signatures(self, doc: Document, directors: list):
        LINE      = "_" * 26
        WIDE_LINE = "_" * 28

        for i in range(0, len(directors), 2):
            pair = directors[i: i + 2]

            p = doc.add_paragraph()
            _set_para_spacing(p, before_pt=6, after_pt=0)
            if len(pair) >= 1:
                _add_run(p, LINE, size_pt=9.5)
            if len(pair) >= 2:
                _add_run(p, "\t\t\t\t", size_pt=9.5)
                _add_run(p, WIDE_LINE, size_pt=9.5)

            p = doc.add_paragraph()
            _set_para_spacing(p, after_pt=4)
            if len(pair) >= 1:
                _add_run(p, pair[0], size_pt=9.5)
            if len(pair) >= 2:
                _add_run(p, "\t\t\t\t\t", size_pt=9.5)
                _add_run(p, pair[1], size_pt=9.5)

            p = doc.add_paragraph()
            _set_para_spacing(p, after_pt=0)
